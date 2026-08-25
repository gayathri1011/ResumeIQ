"""Generate sample resume files for tests."""

from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document


SAMPLE_RESUME_TEXT = {
    "name": "Jane Doe",
    "email": "jane.doe@example.com",
    "phone": "+1 555-010-2000",
    "summary": "Software engineer with 5 years of experience building web applications.",
    "experience_title": "Senior Software Engineer",
    "experience_org": "Acme Corp",
    "experience_dates": "Jan 2020 - Present",
    "experience_bullet": "Built REST APIs serving 1M requests per day.",
    "education": "B.S. Computer Science, State University, 2018",
    "skills": "Python, TypeScript, PostgreSQL, FastAPI, React",
}


def create_sample_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    y = 72
    lines = [
        SAMPLE_RESUME_TEXT["name"],
        f"{SAMPLE_RESUME_TEXT['email']} | {SAMPLE_RESUME_TEXT['phone']}",
        "",
        "Professional Summary",
        SAMPLE_RESUME_TEXT["summary"],
        "",
        "Work Experience",
        f"{SAMPLE_RESUME_TEXT['experience_title']} | {SAMPLE_RESUME_TEXT['experience_org']}",
        SAMPLE_RESUME_TEXT["experience_dates"],
        f"• {SAMPLE_RESUME_TEXT['experience_bullet']}",
        "",
        "Education",
        SAMPLE_RESUME_TEXT["education"],
        "",
        "Skills",
        SAMPLE_RESUME_TEXT["skills"],
    ]
    for line in lines:
        page.insert_text((72, y), line, fontsize=11)
        y += 16
    doc.save(path)
    doc.close()


def create_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph(SAMPLE_RESUME_TEXT["name"])
    doc.add_paragraph(f"{SAMPLE_RESUME_TEXT['email']} | {SAMPLE_RESUME_TEXT['phone']}")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(SAMPLE_RESUME_TEXT["summary"])
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph(
        f"{SAMPLE_RESUME_TEXT['experience_title']} | {SAMPLE_RESUME_TEXT['experience_org']}"
    )
    doc.add_paragraph(SAMPLE_RESUME_TEXT["experience_dates"])
    doc.add_paragraph(f"• {SAMPLE_RESUME_TEXT['experience_bullet']}")
    doc.add_heading("Education", level=1)
    doc.add_paragraph(SAMPLE_RESUME_TEXT["education"])
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(SAMPLE_RESUME_TEXT["skills"])
    doc.save(path)


def create_minimal_pdf_no_certifications(path: Path) -> None:
    """Resume deliberately missing certifications section."""
    create_sample_pdf(path)


def create_empty_pdf(path: Path) -> None:
    doc = fitz.open()
    doc.new_page()
    doc.save(path)
    doc.close()


def create_corrupted_file(path: Path) -> None:
    path.write_bytes(b"not a real pdf content")
